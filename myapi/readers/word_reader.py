import os
from langchain.schema import Document
from docx import Document as DocxDocument

class WordLoader:
    def __init__(self, folder_path, extensions=None):
        self.folder_path = folder_path
        self.extensions = extensions or [".docx"]

    def load(self):
        documents = []
        for filename in os.listdir(self.folder_path):
            ext = os.path.splitext(filename)[1].lower()
            if ext in self.extensions:
                filepath = os.path.join(self.folder_path, filename)
                try:
                    doc = DocxDocument(filepath)
                    full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                    if full_text.strip():
                        documents.append(Document(page_content=full_text, metadata={"source": filepath.replace('stagging/','')}))
                        print(f"Loaded Word doc: {filename}")
                    else:
                        print(f"No text found in Word doc: {filename}")
                except Exception as e:
                    print(f"Failed to load Word doc {filename}: {e}")
        return documents
    
